import io
import docx
from typing import Dict, Any
from .base import BaseParser


class DocxParser(BaseParser):
    """DOCX文件解析器"""
    
    async def extract_text(self, file_data: bytes) -> str:
        """
        从DOCX文件中提取文本内容
        
        Args:
            file_data: DOCX文件二进制数据
            
        Returns:
            str: 提取的文本内容
        """
        text = ""
        try:
            docx_file = io.BytesIO(file_data)
            doc = docx.Document(docx_file)
            
            # 提取段落文本
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            # 实际应用中应该使用日志记录错误
            print(f"DOCX解析错误: {str(e)}")
            return ""
    
    async def get_metadata(self, file_data: bytes) -> Dict[str, Any]:
        """
        从DOCX文件中提取元数据
        
        Args:
            file_data: DOCX文件二进制数据
            
        Returns:
            Dict[str, Any]: 元数据字典
        """
        metadata = {}
        try:
            docx_file = io.BytesIO(file_data)
            doc = docx.Document(docx_file)
            
            # 获取文档属性
            core_properties = doc.core_properties
            
            # 提取常见元数据
            if core_properties.author:
                metadata['author'] = core_properties.author
            if core_properties.title:
                metadata['title'] = core_properties.title
            if core_properties.created:
                metadata['created'] = core_properties.created
            if core_properties.modified:
                metadata['modified'] = core_properties.modified
            if core_properties.subject:
                metadata['subject'] = core_properties.subject
            if core_properties.keywords:
                metadata['keywords'] = core_properties.keywords
            
            # 添加统计信息
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            return metadata
        except Exception as e:
            # 实际应用中应该使用日志记录错误
            print(f"DOCX元数据提取错误: {str(e)}")
            return {}
    
    def supported_mime_types(self) -> list:
        """
        获取解析器支持的MIME类型列表
        
        Returns:
            list: 支持的MIME类型列表
        """
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ] 